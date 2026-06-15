#!/usr/bin/env python3
"""
Thesis Formatter — 西安石油大学本科毕业设计(论文)模板
Features: page setup, styles, TOC, section breaks, page breaks,
          headers/footers per section, captions, references.

Usage:
  python thesis_formatter.py "input.docx" ["output.docx"] [--doc-title="标题"] [--no-toc]
"""

import re, shutil, os, sys, argparse, tempfile
from copy import deepcopy

import docx
from docx import Document
from docx.shared import Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree

# ── Constants ──────────────────────────────────────────────
SIMSUN = '宋体'
TNR = 'Times New Roman'
BLACK = RGBColor(0, 0, 0)

# Page dimensions in EMU (A4)
PAGE_W_EMU = 7560310   # 210mm → ~596pt → but we use standard 11906 twips
PAGE_H_EMU = 10692130
# Actually let's use the standard A4 in EMU directly:
# A4 = 210mm x 297mm
# 1mm = 36000 EMU (approximately, but Word uses twips internally)
# Standard Word A4: 11906 x 16838 twips
# 1 twip = 635 EMU
# So: 11906*635 = 7560310, 16838*635 = 10692130 ✓
MARGIN_LEFT_EMU  = int(3.0 * 360000)   # 3.0cm → EMU (using 1cm=360000 EMU)
MARGIN_RIGHT_EMU = int(2.5 * 360000)   # 2.5cm
MARGIN_TOP_EMU   = int(2.5 * 360000)   # 2.5cm
MARGIN_BOT_EMU   = int(2.5 * 360000)   # 2.5cm
# Actually python-docx uses EMU. Cm(3.0) = 3*360000 = 1080000 EMU
# Let's use the Cm() helper for margins and convert to int for XML

# ── Heading detection patterns ─────────────────────────────
# Chinese numeral characters (used across multiple regex patterns)
CN_NUM_CHARS = '一二三四五六七八九十百千'

# Heading 1 patterns: 一、二、三... / 第一章 / 摘要 / ABSTRACT / 参考文献 / etc.
H1_CN = re.compile(rf'^[{CN_NUM_CHARS}]+[、．\.]')     # 一、, 十一．, 二十一、
H1_DI_ZHANG  = re.compile(rf'^第[{CN_NUM_CHARS}\d]+章') # 第一章, 第1章, 第十一章
H1_CN_SPACE  = re.compile(rf'^[{CN_NUM_CHARS}]{{1,3}}\s') # 一 (bare, space after)
H1_NUM_SPACE = re.compile(r'^[1-9]\d*\s(?!\d)')         # 1 Introduction (no dot after number)
H1_NUM = re.compile(r'^[1-9]\d*[\.\、](?!\d)')          # 1. / 2、 (not 1.1/2.3)
SPECIAL_H1 = {
    '摘要', 'ABSTRACT', 'Abstract', '参考文献', '致谢', '谢辞',
    '绪论', '引言', '前言', '结论',
    '目 录', '目录',           # TOC entries
    '鸣谢',                     # alternative thanks
    '附录一', '附录二', '附录三', '附录四', '附录五',
}

# Heading 2 patterns:
#   Standard: 1.1 / 2.3 (需在数字编号后有空格的章节标题)
#   Chinese: （一）（二）（三）...
#   第X节: 第一节, 第二节, 第1节, etc.
H2_NUM = re.compile(r'^[1-9]\d*\.\d+\s')          # 1.1 XXX
H2_CN  = re.compile(r'^（[一二三四五六七八九十]+）')  # （一）XXX
H2_DI_JIE = re.compile(rf'^第[{CN_NUM_CHARS}\d]+节') # 第一节

# Heading 3 patterns:
#   Standard: 1.1.1 / 2.3.4
#   NOTE: （1）（2）style is NOT used as Heading 3 — body paragraphs
#         starting with （N）are typically enumeration points, not headings.
H3_NUM = re.compile(r'^[1-9]\d*\.\d+\.\d+\s')     # 1.1.1 XXX

# Headings that trigger a section break before them (legacy, now mode-dependent)
SECTION_BREAK_SPECIAL = {'参考文献', '致谢', '谢辞'}

# Chinese numeral → Arabic conversion
CN_DIGITS = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
             '六': 6, '七': 7, '八': 8, '九': 9}

def cn_num_to_arabic(cn):
    """Convert Chinese numeral to Arabic integer.
    Supports: 一-九, 十-十九, 二十-九十九
    Returns None if conversion fails."""
    if cn in CN_DIGITS:
        return CN_DIGITS[cn]
    if cn == '十':
        return 10
    # X十Y format: 二十一 → 21
    if len(cn) == 3 and cn[1] == '十':
        tens = CN_DIGITS.get(cn[0], 0)
        ones = CN_DIGITS.get(cn[2], 0)
        if tens > 0 and ones > 0:
            return tens * 10 + ones
    # 十Y format: 十一 → 11
    if len(cn) == 2 and cn[0] == '十':
        ones = CN_DIGITS.get(cn[1], 0)
        if ones > 0:
            return 10 + ones
    # X十 format: 二十 → 20
    if len(cn) == 2 and cn[1] == '十':
        tens = CN_DIGITS.get(cn[0], 0)
        if tens > 0:
            return tens * 10
    return None


def rename_h1_cn_to_arabic(para):
    """Rename a Heading 1 from Chinese numeral to Arabic.
    E.g., '一、绪论' → '1. 绪论', '十一、XXX' → '11. XXX'
    Returns True if renamed."""
    text = para.text.strip()
    m = H1_CN.match(text)
    if not m:
        return False
    cn_prefix = m.group()           # e.g., '一、' or '十一．'
    cn_numeral = cn_prefix[:-1]     # e.g., '一' or '十一'
    arabic = cn_num_to_arabic(cn_numeral)
    if arabic is None:
        return False
    new_text = f'{arabic}. {text[len(cn_prefix):].strip()}'
    # Clear existing runs and replace with single renamed run
    for run_elem in list(para._element.findall(qn('w:r'))):
        para._element.remove(run_elem)
    new_run = para.add_run(new_text)
    set_run_xml(new_run, 15, bold=True)
    return True


# ── XML/Run helpers ────────────────────────────────────────
def set_run_xml(run, size_pt, bold, ea=SIMSUN, latin=TNR):
    """Set font properties on a run using XML for east-Asian font support."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s />' % nsdecls("w"))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), ea)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:cs'), latin)

    for tag in [qn('w:sz'), qn('w:szCs')]:
        el = rPr.find(tag)
        if el is None:
            el = parse_xml('<w:%s %s w:val="%d"/>' % (
                tag.split('}')[1], nsdecls("w"), int(size_pt * 2)))
            rPr.append(el)
        else:
            el.set(qn('w:val'), str(int(size_pt * 2)))

    b = rPr.find(qn('w:b'))
    if bold:
        if b is None:
            rPr.append(parse_xml('<w:b %s/>' % nsdecls("w")))
    else:
        if b is not None:
            rPr.remove(b)

    c = rPr.find(qn('w:color'))
    if c is None:
        c = parse_xml('<w:color %s w:val="000000"/>' % nsdecls("w"))
        rPr.append(c)

    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = BLACK


def config_style(style, size, bold, alignment, first_indent=None,
                 line_spacing=None, space_before=None, space_after=None,
                 left_indent=None):
    """Configure a Word style with Chinese font support."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts %s />' % nsdecls("w"))
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), SIMSUN)
    rFonts.set(qn('w:ascii'), TNR)
    rFonts.set(qn('w:hAnsi'), TNR)
    style.font.name = TNR
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK
    pf = style.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if left_indent is not None:
        pf.left_indent = left_indent


# ── Markdown cleanup ─────────────────────────────────────────
# Patterns for stripping markdown artifacts from docx text
MD_HEADING = re.compile(r'^(#{1,3})\s+')
MD_BOLD = re.compile(r'\*\*(.+?)\*\*')
MD_ITALIC = re.compile(r'\*(.+?)\*')
MD_BOLD_ALT = re.compile(r'__(.+?)__')       # alternate bold (AI-generated)
MD_STRIKE = re.compile(r'~~(.+?)~~')          # strikethrough


def strip_md(text):
    """Strip markdown decorators for text analysis (no document modification)."""
    s = text.strip()
    s = MD_HEADING.sub('', s).strip()
    s = MD_BOLD.sub(r'\1', s)
    s = MD_ITALIC.sub(r'\1', s)
    s = MD_BOLD_ALT.sub(r'\1', s)
    s = MD_STRIKE.sub(r'\1', s)
    return s.strip()


def clean_document_markdown(doc):
    """Remove markdown artifacts from ALL paragraphs in the document.
    - Strips ### / ## / # heading prefixes from first text run
    - Removes **bold**, *italic*, __bold_alt__, and ~~strikethrough~~ markers from all runs
    Returns number of modified paragraphs."""
    cleaned = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # Only process paragraphs that actually have markdown artifacts
        if not ('#' in text[:6] or '**' in text or '*' in text
                or '__' in text or '~~' in text):
            continue

        modified = False
        runs = p._element.findall(qn('w:r'))

        # Phase 1: Strip heading marker from first text-containing run
        for run_elem in runs:
            t_elems = run_elem.findall(qn('w:t'))
            for t in t_elems:
                if t.text:
                    m = MD_HEADING.match(t.text)
                    if m:
                        t.text = t.text[m.end():]
                        modified = True
                    break  # Only strip from first text element

        # Phase 2: Strip **, *, __, and ~~ markers from all runs
        for run_elem in runs:
            t_elems = run_elem.findall(qn('w:t'))
            for t in t_elems:
                if t.text and ('**' in t.text or '*' in t.text
                               or '__' in t.text or '~~' in t.text):
                    t.text = MD_BOLD.sub(r'\1', t.text)
                    t.text = MD_ITALIC.sub(r'\1', t.text)
                    t.text = MD_BOLD_ALT.sub(r'\1', t.text)
                    t.text = MD_STRIKE.sub(r'\1', t.text)
                    modified = True

        if modified:
            cleaned += 1

    return cleaned


# ── Heading identification ──────────────────────────────────
def identify(text):
    """Identify heading level from paragraph text.
    Strips markdown decorators before pattern matching.
    Detection order (H3 → H2 → H1):
      Heading 1: 第一章/第二章... / 一、二、三... / 1. / 2. /
                 摘要/ABSTRACT/参考文献/致谢/绪论/结论/附录X/目录
      Heading 2: （一）（二）... / 第一节... / 1.1 / 2.3 (数字编号)
      Heading 3: 1.1.1 / 2.3.4 (数字编号)
    """
    s = strip_md(text)  # Strip markdown (#, **, *) before matching
    if not s:
        return None

    # Level 3 (check first, most specific)
    if H3_NUM.match(s):
        return 'Heading 3'

    # Level 2
    if H2_NUM.match(s) or H2_CN.match(s):
        return 'Heading 2'
    if H2_DI_JIE.match(s):          # 第一节, 第1节
        return 'Heading 2'

    # Level 1 — check most specific patterns first
    if H1_DI_ZHANG.match(s):        # 第一章, 第1章
        return 'Heading 1'
    if H1_CN.match(s) or H1_NUM.match(s):
        return 'Heading 1'
    if H1_CN_SPACE.match(s):        # 一 (bare, space after)
        return 'Heading 1'
    if H1_NUM_SPACE.match(s):       # 1 (bare, space after)
        return 'Heading 1'
    if s.startswith('附录'):
        return 'Heading 1'
    for kw in SPECIAL_H1:
        if s == kw or s.startswith(kw):
            return 'Heading 1'

    return None


# ── Formatting-based heading inference ──────────────────────
def infer_heading_from_formatting(paragraph):
    """Fallback heading detection based on run-level formatting cues.

    Examines the first run's font size (w:sz in half-points), bold state (w:b),
    and paragraph alignment (w:jc). Returns 'Heading 1', 'Heading 2', 'Heading 3',
    or None if formatting is ambiguous.

    Heuristic rules (calibrated for Chinese academic conventions):
      - Centered + bold + sz >= 30 (15pt)  → Heading 1 (小三)
      - Bold + sz >= 28 (14pt)             → Heading 2 (四号)
      - Bold + sz >= 24 (12pt)             → Heading 3 (小四)
      - Rejects paragraphs longer than 80 chars (body text, not a heading)
      - Rejects paragraphs that are not bold (headings are uniformly bold)
    """
    text = paragraph.text.strip()
    if not text or len(text) > 80:
        return None

    pPr = paragraph._element.find(qn('w:pPr'))
    alignment = None
    if pPr is not None:
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            alignment = jc.get(qn('w:val'))

    runs = paragraph._element.findall(qn('w:r'))
    if not runs:
        return None

    first_run = runs[0]
    rPr = first_run.find(qn('w:rPr'))
    if rPr is None:
        return None

    sz_elem = rPr.find(qn('w:sz'))
    if sz_elem is None:
        return None

    try:
        sz = int(sz_elem.get(qn('w:val')))  # half-points
    except (ValueError, TypeError):
        return None

    is_bold = rPr.find(qn('w:b')) is not None
    if not is_bold:
        return None

    # Centered + bold + 15pt+ → Heading 1 (小三)
    if alignment == 'center' and sz >= 30:
        return 'Heading 1'
    # Bold + 14pt+ → Heading 2 (四号)
    if sz >= 28:
        return 'Heading 2'
    # Bold + 12pt+ → Heading 3 (小四)
    if sz >= 24:
        return 'Heading 3'

    return None


def identify_paragraph(paragraph):
    """Identify heading level using regex patterns first, then formatting fallback.
    This should be used instead of identify(text) when the paragraph object
    is available, so that formatting-based inference can act as a fallback."""
    text = paragraph.text.strip()
    result = identify(text)  # regex-based, uses strip_md internally
    if result:
        return result
    return infer_heading_from_formatting(paragraph)


# ── Section break (XML manipulation) ────────────────────────
def make_sectPr_xml(page_w_emu, page_h_emu, left_emu, right_emu,
                    top_emu, bottom_emu):
    """Create a <w:sectPr> element with page dimensions and margins.
    All values in EMU."""
    sectPr = OxmlElement('w:sectPr')

    # Page size
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), str(page_w_emu))
    pgSz.set(qn('w:h'), str(page_h_emu))
    sectPr.append(pgSz)

    # Page margins
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:left'), str(left_emu))
    pgMar.set(qn('w:right'), str(right_emu))
    pgMar.set(qn('w:top'), str(top_emu))
    pgMar.set(qn('w:bottom'), str(bottom_emu))
    pgMar.set(qn('w:header'), '851')   # default header distance ~1.5cm
    pgMar.set(qn('w:footer'), '851')   # default footer distance
    pgMar.set(qn('w:gutter'), '0')
    sectPr.append(pgMar)

    # Columns (single)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:space'), '720')
    sectPr.append(cols)

    # Document grid
    docGrid = OxmlElement('w:docGrid')
    docGrid.set(qn('w:linePitch'), '360')
    sectPr.append(docGrid)

    return sectPr


def add_section_break_at(para):
    """Make this paragraph the LAST paragraph of a section.
    Adds <w:sectPr> to the paragraph's <w:pPr> with proper page dimensions.
    The NEXT paragraph after this one will start a new section."""
    pPr = para._element.get_or_add_pPr()

    # Remove any existing sectPr in this paragraph
    for old in pPr.findall(qn('w:sectPr')):
        pPr.remove(old)

    sectPr = make_sectPr_xml(PAGE_W_EMU, PAGE_H_EMU,
                             MARGIN_LEFT_EMU, MARGIN_RIGHT_EMU,
                             MARGIN_TOP_EMU, MARGIN_BOT_EMU)
    pPr.append(sectPr)


def clear_all_section_breaks(doc):
    """Remove all existing <w:sectPr> from paragraph bodies.
    This cleans up leftover section breaks from previous formatting runs,
    so only our own section breaks remain."""
    removed = 0
    for p in doc.paragraphs:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is None:
            continue
        for old in pPr.findall(qn('w:sectPr')):
            pPr.remove(old)
            removed += 1
    return removed


# ── Page break ──────────────────────────────────────────────
def add_page_break_before(para):
    """Set page break before this paragraph (used for Heading 1 chapters)."""
    para.paragraph_format.page_break_before = True


# ── Header/Footer helpers ──────────────────────────────────
def clear_header_footer(section):
    """Remove all paragraphs from header and footer of a section."""
    for attr in ['header', 'footer', 'first_page_header', 'first_page_footer',
                 'even_page_header', 'even_page_footer']:
        try:
            part = getattr(section, attr)
            if part is None:
                continue
            # Remove all paragraph elements
            for p_elem in list(part._element):
                if p_elem.tag == qn('w:p'):
                    part._element.remove(p_elem)
        except Exception:
            pass


def set_section_header(section, text, size_pt=10.5):
    """Set centered header with given text."""
    try:
        header = section.header
        if header is None:
            return
        header.is_linked_to_previous = False
        # Clear existing
        for p_elem in list(header._element):
            if p_elem.tag == qn('w:p'):
                header._element.remove(p_elem)
        # Add header paragraph
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.clear()
        run = hp.add_run(text)
        set_run_xml(run, size_pt, bold=False)
    except Exception as e:
        print(f"  [WARN] Header setup failed: {e}")


def set_section_footer_page_number(section, size_pt=10.5):
    """Set centered footer with PAGE field (page number)."""
    try:
        footer = section.footer
        if footer is None:
            return
        footer.is_linked_to_previous = False
        # Clear existing
        for p_elem in list(footer._element):
            if p_elem.tag == qn('w:p'):
                footer._element.remove(p_elem)
        # Add footer paragraph
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        run = fp.add_run()
        set_run_xml(run, size_pt, bold=False, latin=TNR, ea=SIMSUN)
        # Insert PAGE field
        run._element.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls("w")))
        run._element.append(parse_xml('<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls("w")))
        run._element.append(parse_xml('<w:fldChar %s w:fldCharType="separate"/>' % nsdecls("w")))
        run._element.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls("w")))
    except Exception as e:
        print(f"  [WARN] Footer setup failed: {e}")


def set_page_number_restart(section, start=1):
    """Set page numbering to restart at `start` for this section."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.insert(0, pgNumType)
    pgNumType.set(qn('w:start'), str(start))


def set_page_number_format(section, fmt='decimal', start=None):
    """Set page number format (decimal, upperRoman, lowerRoman, etc.)
    and optionally restart page numbering.

    Args:
        fmt: 'decimal', 'upperRoman', 'lowerRoman', 'upperLetter', etc.
        start: If provided, restart page numbering at this value.
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            sectPr.insert(list(sectPr).index(pgSz), pgNumType)
        else:
            sectPr.insert(0, pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))


def enable_even_odd_headers(section):
    """Enable different odd/even page headers for a section."""
    sectPr = section._sectPr
    if sectPr.find(qn('w:evenAndOddHeaders')) is None:
        even_odd = OxmlElement('w:evenAndOddHeaders')
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            sectPr.insert(list(sectPr).index(pgSz), even_odd)
        else:
            sectPr.insert(0, even_odd)


def set_section_even_header(section, text, size_pt=10.5):
    """Set centered header for even pages only.
    Requires enable_even_odd_headers() to have been called on the section."""
    try:
        header = section.even_page_header
        if header is None:
            return
        header.is_linked_to_previous = False
        for p_elem in list(header._element):
            if p_elem.tag == qn('w:p'):
                header._element.remove(p_elem)
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.clear()
        run = hp.add_run(text)
        set_run_xml(run, size_pt, bold=False)
    except Exception as e:
        print(f"  [WARN] Even header setup failed: {e}")


# ── Main pipeline ───────────────────────────────────────────
def format_thesis(input_path, output_path, doc_title="论文", no_toc=False,
                  mode="xsyu_thesis", course_title=None):
    """Main formatting pipeline.

    Args:
        input_path: Path to input .docx file
        output_path: Path to output .docx file
        doc_title: Document title (used as header text in XSYU mode)
        no_toc: Skip TOC insertion message
        mode: "xsyu_thesis" (XSYU formal thesis) or "homework" (daily assignment)
        course_title: Header text for homework mode (falls back to doc_title if None)
    """

    # ── Step 0: Backup existing output ──
    if os.path.exists(output_path):
        backup = output_path.replace('.docx', '.backup.docx')
        if not backup.endswith('.backup.docx'):
            backup = output_path + '.backup.docx'
        shutil.copy2(output_path, backup)
        print(f"[0] Backup: {backup}")

    # ── Step 1: Load & Analyze ──
    print(f"[1] Loading: {input_path}")
    doc = Document(input_path)

    # ── Step 1b: Clean markdown artifacts ──
    md_cleaned = clean_document_markdown(doc)
    if md_cleaned:
        print(f"[1b] Markdown cleanup: {md_cleaned} paragraphs (###/**, * stripped)")

    # Identify headings and structure
    heading_map = {}   # para_idx → style_name
    cover_end = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        h = identify_paragraph(p)
        if h:
            heading_map[i] = h
            if cover_end == 0:
                cover_end = i  # First heading = end of cover

    # Find special headings
    ref_idx = None   # 参考文献
    thanks_idx = None  # 致谢
    appendix_indices = []  # 附录
    for idx in sorted(heading_map):
        t = doc.paragraphs[idx].text.strip()
        if heading_map[idx] == 'Heading 1':
            if t == '参考文献':
                ref_idx = idx
            elif t in ('致谢', '谢辞'):
                thanks_idx = idx
            elif t.startswith('附录'):
                appendix_indices.append(idx)

    print(f"[2] Structure: {len(heading_map)} headings, "
          f"cover ends at P{cover_end}, ref=P{ref_idx}, "
          f"thanks=P{thanks_idx}, appendices={appendix_indices}")

    # ── Step 2: Apply styles ──
    config_style(doc.styles['Normal'], 12, False, WD_ALIGN_PARAGRAPH.JUSTIFY,
                 first_indent=Pt(24), line_spacing=1.25,
                 space_before=Pt(0), space_after=Pt(0))
    config_style(doc.styles['Heading 1'], 15, True, WD_ALIGN_PARAGRAPH.CENTER,
                 first_indent=Pt(0), line_spacing=1.25,
                 space_before=Pt(12), space_after=Pt(24))
    config_style(doc.styles['Heading 2'], 14, True, WD_ALIGN_PARAGRAPH.LEFT,
                 first_indent=Pt(0), line_spacing=1.25,
                 space_before=Pt(6), space_after=Pt(6))
    config_style(doc.styles['Heading 3'], 12, True, WD_ALIGN_PARAGRAPH.LEFT,
                 first_indent=Pt(0), line_spacing=1.25,
                 space_before=Pt(6), space_after=Pt(6))
    config_style(doc.styles['List Paragraph'], 10.5, False,
                 WD_ALIGN_PARAGRAPH.JUSTIFY,
                 first_indent=Pt(0), left_indent=Pt(0), line_spacing=1.25,
                 space_before=Pt(0), space_after=Pt(0))
    try:
        config_style(doc.styles['toc 1'], 12, False, WD_ALIGN_PARAGRAPH.LEFT,
                     first_indent=Pt(0), line_spacing=1.2)
        config_style(doc.styles['toc 2'], 12, False, WD_ALIGN_PARAGRAPH.LEFT,
                     first_indent=Pt(0), left_indent=Pt(24))
        config_style(doc.styles['toc 3'], 12, False, WD_ALIGN_PARAGRAPH.LEFT,
                     first_indent=Pt(0), left_indent=Pt(48))
    except Exception:
        pass
    print("[3] Styles OK")

    # ── Step 4: Format cover page (title, abstract, keywords) ──
    # (Must run BEFORE section/page breaks because it may insert paragraphs)
    # 西安石油大学模板规范:
    #   论文题目: 宋体, 小三(15pt), 加粗, 居中, 1.25x行距, 段前0.5行, 段后1行
    #   摘要标题: 宋体, 小三(15pt), 加粗, 居中, 1.25x行距, 段前1行, 段后2行
    #   摘要正文: (Normal 保持)
    #   关键词: 宋体, 小四(12pt), 顶格左对齐, "关键词："加粗, 分号分隔

    cover_title_idx = None
    abstract_idx = None    # paragraph containing "摘要" (may be combined with body)
    abstract_body_idx = None  # split-off abstract body (if split)
    keywords_idx = None

    # Find cover boundary: the first paragraph that is a "real" heading
    # (not 摘要/ABSTRACT, i.e. a chapter heading like 一、绪论)
    cover_boundary = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        h = identify_paragraph(p)
        if h == 'Heading 1':
            # 摘要/ABSTRACT are part of the cover, not a chapter
            is_abstract = (
                text == '摘要' or text.startswith('摘要') or text.startswith('摘　要')
                or text == 'ABSTRACT' or text == 'Abstract'
            )
            if not is_abstract:
                cover_boundary = i
                break

    # Scan cover area for title, abstract, keywords (up to cover_boundary)
    for i in range(cover_boundary):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        if cover_title_idx is None:
            cover_title_idx = i
            continue
        # Detect abstract: starts with 摘要 or 摘 要 (strip_md for safety)
        clean = strip_md(text)
        if abstract_idx is None and ((clean.startswith('摘') and ('要' in clean[:4])) or clean == '摘要'):
            abstract_idx = i
            continue
        # Detect keywords (strip_md for safety)
        if keywords_idx is None and (clean.startswith('关键词') or clean.startswith('關鍵詞')):
            keywords_idx = i
            continue

    # 4a: Format cover title
    if cover_title_idx is not None:
        p = doc.paragraphs[cover_title_idx]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_before = Pt(6)    # 段前0.5行 ≈ 6pt
        p.paragraph_format.space_after = Pt(12)     # 段后1行 ≈ 12pt
        for run in p.runs:
            set_run_xml(run, 15, bold=True)          # 小三 15pt
        print(f"[4a] Cover title: P{cover_title_idx}")

        # 4b: Process abstract — split "摘要" heading from body if they are in same paragraph
    if abstract_idx is not None:
        p = doc.paragraphs[abstract_idx]
        text = p.text
        heading_text = None
        body_text = None

        # 1. 先尝试冒号分隔符（保持原有逻辑）
        for delimiter in ['摘　要：', '摘　要:', '摘要：', '摘要:', '摘 要：', '摘 要:']:
            if text.startswith(delimiter):
                heading_text = '摘要'
                body_text = text[len(delimiter):].strip()
                break

        # 2. 如果没有冒号，但以“摘要”开头且其后有内容（包括空格、换行等）
        if heading_text is None and (text.startswith('摘要') or text.startswith('摘　要')):
            # 提取“摘要”二字作为标题，剩余部分作为正文
            # 注意：可能“摘要”后面紧跟内容，无任何分隔符
            match = re.match(r'^(摘要|摘　要)(.*)$', text, re.DOTALL)
            if match:
                heading_text = '摘要'
                body_text = match.group(2).strip()
                # 如果正文为空，则不拆分（保持原样）
                if not body_text:
                    heading_text = None

        if heading_text and body_text:
            # 创建新段落作为摘要标题（插入在当前段落之前）
            new_p_elem = OxmlElement('w:p')
            p._element.addprevious(new_p_elem)
            # 构建标题段落
            new_p_pr = OxmlElement('w:pPr')
            new_p_elem.append(new_p_pr)
            new_r = OxmlElement('w:r')
            new_rPr = OxmlElement('w:rPr')
            new_r.append(new_rPr)
            new_t = OxmlElement('w:t')
            new_t.text = heading_text
            new_t.set(qn('xml:space'), 'preserve')
            new_r.append(new_t)
            new_p_elem.append(new_r)

            # 清空原段落并填充正文内容
            for run_elem in list(p._element.findall(qn('w:r'))):
                p._element.remove(run_elem)
            body_r = OxmlElement('w:r')
            body_rPr = OxmlElement('w:rPr')
            body_r.append(body_rPr)
            body_t = OxmlElement('w:t')
            body_t.text = body_text
            body_t.set(qn('xml:space'), 'preserve')
            body_r.append(body_t)
            p._element.append(body_r)

            # 设置标题段落的格式（居中，Heading 1 样式）
            pPr = new_p_elem.find(qn('w:pPr'))
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            pPr.append(jc)
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '240')
            spacing.set(qn('w:after'), '480')
            spacing.set(qn('w:line'), '300')
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '0')
            pPr.append(ind)

            # 标题运行字体
            run_rPr = new_r.find(qn('w:rPr'))
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), SIMSUN)
            rFonts.set(qn('w:ascii'), TNR)
            rFonts.set(qn('w:hAnsi'), TNR)
            run_rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '30')
            run_rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '30')
            run_rPr.append(szCs)
            b = OxmlElement('w:b')
            run_rPr.append(b)

            # 设置正文段落格式（Normal 样式，首行缩进2字符）
            body_pPr = p._element.get_or_add_pPr()
            for old in body_pPr.findall(qn('w:spacing')):
                body_pPr.remove(old)
            for old in body_pPr.findall(qn('w:ind')):
                body_pPr.remove(old)
            body_spacing = OxmlElement('w:spacing')
            body_spacing.set(qn('w:line'), '300')
            body_spacing.set(qn('w:lineRule'), 'auto')
            body_pPr.append(body_spacing)
            body_ind = OxmlElement('w:ind')
            body_ind.set(qn('w:firstLine'), '480')
            body_pPr.append(body_ind)

            body_run_rPr = body_r.find(qn('w:rPr'))
            body_rFonts = OxmlElement('w:rFonts')
            body_rFonts.set(qn('w:eastAsia'), SIMSUN)
            body_rFonts.set(qn('w:ascii'), TNR)
            body_rFonts.set(qn('w:hAnsi'), TNR)
            body_run_rPr.append(body_rFonts)
            body_sz = OxmlElement('w:sz')
            body_sz.set(qn('w:val'), '24')
            body_run_rPr.append(body_sz)

            # 记录拆分后的正文段落索引（原段落现在变成正文）
            abstract_body_idx = abstract_idx + 1
            if keywords_idx is not None and keywords_idx > abstract_idx:
                keywords_idx += 1

            # 将新标题段落标记为 Heading 1
            new_p_style = new_p_elem.get_or_add_pPr()
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), 'Heading1')
            new_p_style.append(pStyle)
            heading_map[abstract_idx] = 'Heading 1'  # 原段落现在是正文，但 map 中仍保留标题位置？稍后会重新扫描，不影响

            print(f"[4b] Abstract split: '{heading_text}' heading + body ({len(body_text)} chars), "
                  f"keywords_idx adjusted to P{keywords_idx}")
        else:
            # 没有拆分（例如只有“摘要”二字，无正文，或无法识别）
            p.style = doc.styles['Heading 1']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(24)
            for run in p.runs:
                set_run_xml(run, 15, bold=True)
            heading_map[abstract_idx] = 'Heading 1'
            print(f"[4b] Abstract heading formatted (no split): P{abstract_idx}")

    # 4c: Add blank line between abstract body and keywords
    if abstract_body_idx is not None and keywords_idx is not None:
        empty_p = OxmlElement('w:p')
        doc.paragraphs[keywords_idx]._element.addprevious(empty_p)
        keywords_idx += 1  # keywords shifted by blank line insert
        print(f"[4c] Blank line inserted, keywords shifted to P{keywords_idx}")

    # 6d: Format keywords
    if keywords_idx is not None:
        p = doc.paragraphs[keywords_idx]
        text = p.text.strip()

        # Parse: "关键词：xxx；xxx" → bold "关键词：" + normal "xxx；xxx"
        kw_label = None
        kw_content = None
        for delim in ['关键词：', '關鍵詞：', '关键词:', '關鍵詞:']:
            if text.startswith(delim):
                kw_label = delim
                kw_content = text[len(delim):].strip()
                break

        if kw_label and kw_content:
            # Clear runs and rebuild
            for run_elem in list(p._element.findall(qn('w:r'))):
                p._element.remove(run_elem)

            # Bold "关键词："
            bold_r = OxmlElement('w:r')
            bold_rPr = OxmlElement('w:rPr')
            bold_rFonts = OxmlElement('w:rFonts')
            bold_rFonts.set(qn('w:eastAsia'), SIMSUN)
            bold_rFonts.set(qn('w:ascii'), TNR)
            bold_rFonts.set(qn('w:hAnsi'), TNR)
            bold_rPr.append(bold_rFonts)
            bold_sz = OxmlElement('w:sz')
            bold_sz.set(qn('w:val'), '24')  # 12pt 小四
            bold_rPr.append(bold_sz)
            bold_b = OxmlElement('w:b')
            bold_rPr.append(bold_b)
            bold_t = OxmlElement('w:t')
            bold_t.text = kw_label
            bold_t.set(qn('xml:space'), 'preserve')
            bold_r.append(bold_rPr)
            bold_r.append(bold_t)
            p._element.append(bold_r)

            # Normal keywords content
            norm_r = OxmlElement('w:r')
            norm_rPr = OxmlElement('w:rPr')
            norm_rFonts = OxmlElement('w:rFonts')
            norm_rFonts.set(qn('w:eastAsia'), SIMSUN)
            norm_rFonts.set(qn('w:ascii'), TNR)
            norm_rFonts.set(qn('w:hAnsi'), TNR)
            norm_rPr.append(norm_rFonts)
            norm_sz = OxmlElement('w:sz')
            norm_sz.set(qn('w:val'), '24')  # 12pt
            norm_rPr.append(norm_sz)
            norm_t = OxmlElement('w:t')
            norm_t.text = kw_content
            norm_t.set(qn('xml:space'), 'preserve')
            norm_r.append(norm_rPr)
            norm_r.append(norm_t)
            p._element.append(norm_r)

        # Format paragraph properties: 宋体小四, 顶格左对齐, 1.25x行距
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        print(f"[4d] Keywords formatted: P{keywords_idx}")

    # ── Step 4e: Insert TOC placeholder after keywords (XSYU mode only) ──
    # Creates a blank paragraph that serves as the TOC section between
    # the abstract section (no page numbers) and body chapters.
    toc_placeholder_idx = None
    if mode == 'xsyu_thesis' and keywords_idx is not None:
        toc_p = OxmlElement('w:p')
        # Add an empty run so the paragraph exists but has no visible text
        toc_r = OxmlElement('w:r')
        toc_p.append(toc_r)
        doc.paragraphs[keywords_idx]._element.addnext(toc_p)
        toc_placeholder_idx = keywords_idx + 1
        # Shift indices for affected elements
        if abstract_body_idx is not None and abstract_body_idx > keywords_idx:
            abstract_body_idx += 1
        print(f"[4e] TOC placeholder inserted at P{toc_placeholder_idx} (XSYU mode)")

    # ── Step 4f: Rename Chinese numeral H1 headings to Arabic ──
    #  一、XXX → 1. XXX,  二、XXX → 2. XXX, etc.
    h1_renamed = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        h = identify_paragraph(p)
        if h == 'Heading 1' and H1_CN.match(text):
            # Skip special headings (摘要, 参考文献, 致谢, 附录, etc.)
            if text in SPECIAL_H1 or text.startswith('附录') \
                    or text.startswith('摘要') or text.startswith('摘　要'):
                continue
            if rename_h1_cn_to_arabic(p):
                h1_renamed += 1
    print(f"[4e] H1 CN→Arabic renamed: {h1_renamed} headings")

    # ── Re-scan document structure after cover modifications ──
    # (Paragraph insertions may have shifted indices)
    heading_map = {}
    cover_end = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        h = identify_paragraph(p)
        if h:
            heading_map[i] = h
            if cover_end == 0:
                cover_end = i

    ref_idx = None
    thanks_idx = None
    appendix_indices = []
    for idx in sorted(heading_map):
        t = doc.paragraphs[idx].text.strip()
        if heading_map[idx] == 'Heading 1':
            if t == '参考文献':
                ref_idx = idx
            elif t in ('致谢', '谢辞'):
                thanks_idx = idx
            elif t.startswith('附录'):
                appendix_indices.append(idx)

    # Fix cover_end: use first real chapter heading, not front matter headings
    # Front matter includes: 摘要, ABSTRACT, 目录, 鸣谢
    # This keeps title, abstract, keywords, and TOC on the same cover page
    real_cover_end = cover_end
    for idx in sorted(heading_map):
        t = doc.paragraphs[idx].text.strip()
        if heading_map[idx] == 'Heading 1':
            is_front = (
                t in ('摘要', 'ABSTRACT', 'Abstract', '目录', '目 录', '鸣谢')
                or t.startswith('摘要') or t.startswith('摘　要')
                or t.startswith('目录') or t.startswith('目　录')
            )
            if not is_front:
                real_cover_end = idx
                break
    cover_end = real_cover_end

    # Recalculate section break positions with correct cover_end
    section_break_paras = set()
    section_plan = []  # Built below, maps section index → info dict

    # ── Helper: find the last non-empty paragraph before a given index ──
    def find_prev_para(idx):
        pre = idx - 1
        while pre >= 0 and not doc.paragraphs[pre].text.strip():
            pre -= 1
        return pre

    if mode == 'xsyu_thesis':
        # XSYU thesis mode section structure:
        #   Section 0: Cover + Abstract + Keywords (NO page numbers)
        #   Section 1: TOC placeholder (Roman numerals, odd="目录" / even=school name)
        #   Section 2+: Each chapter in its own section (Arabic, odd/even headers)
        #
        # Break 1: After keywords → separates abstract from TOC
        # Break 2: At TOC placeholder → separates TOC from body
        # Break 3+: Before each non-front-matter H1 → separates chapters

        section_plan = [{'title': '封面与摘要', 'section_type': 'abstract'}]

        # Break 1: After keywords (abstract section end)
        if keywords_idx is not None and keywords_idx > 0:
            section_break_paras.add(keywords_idx)
        elif toc_placeholder_idx is not None and toc_placeholder_idx > 0:
            pre = find_prev_para(toc_placeholder_idx)
            if pre >= 0:
                section_break_paras.add(pre)

        # Break 2: At TOC placeholder (TOC section end)
        if toc_placeholder_idx is not None:
            section_break_paras.add(toc_placeholder_idx)
            section_plan.append({'title': '目录', 'section_type': 'toc'})

        # Break 3+: Before each non-front-matter H1
        for idx in sorted(heading_map):
            if heading_map[idx] != 'Heading 1':
                continue
            t = doc.paragraphs[idx].text.strip()
            is_front_matter = (
                t == '摘要' or t.startswith('摘要') or t.startswith('摘　要')
                or t == 'ABSTRACT' or t == 'Abstract'
                or t == '目录' or t == '目 录' or t.startswith('目录')
                or t == '鸣谢'
            )
            if is_front_matter:
                continue
            pre = find_prev_para(idx)
            if pre >= 0 and pre not in section_break_paras:
                section_break_paras.add(pre)
                section_plan.append({
                    'title': t,
                    'section_type': 'chapter'
                })

    elif mode == 'homework':
        # Homework mode section structure:
        #   Section 0: Cover + Abstract + Keywords (NO page numbers)
        #   Section 1: Body chapters + 致谢 + 附录 (header=course, footer=PAGE from 1)
        #   Section 2: References (header=course, footer=PAGE continuous)
        #
        # Break 1: After keywords (abstract → body)
        # Break 2: Before 参考文献 (body → references)

        section_plan = [{'title': '封面与摘要', 'section_type': 'abstract'}]

        # Break 1: After keywords
        if keywords_idx is not None and keywords_idx > 0:
            section_break_paras.add(keywords_idx)
        section_plan.append({'title': course_title or doc_title, 'section_type': 'body'})

        # Break 2: Before 参考文献
        if ref_idx is not None and ref_idx > 0:
            pre = find_prev_para(ref_idx)
            if pre >= 0 and pre not in section_break_paras:
                section_break_paras.add(pre)
                section_plan.append({'title': '参考文献', 'section_type': 'references'})

    # Apply Heading styles to cover-area headings (before cover_end)
    for idx in sorted(heading_map):
        if idx < cover_end and heading_map[idx] == 'Heading 1':
            doc.paragraphs[idx].style = doc.styles['Heading 1']
            for run in doc.paragraphs[idx].runs:
                set_run_xml(run, 15, bold=True)

    print(f"[5] Re-scanned: {len(heading_map)} headings, "
          f"cover_end=P{cover_end}, ref=P{ref_idx}, "
          f"section_breaks={sorted(section_break_paras)}, "
          f"mode={mode}")

    # ── Step 5b: Clear old section breaks from previous formatting runs ──
    old_breaks = clear_all_section_breaks(doc)
    if old_breaks:
        print(f"[5b] Cleared {old_breaks} old section breaks")

    # ── Step 6: Insert section breaks (XML manipulation) ──
    for break_idx in sorted(section_break_paras):
        para = doc.paragraphs[break_idx]
        add_section_break_at(para)
    print(f"[6] Section breaks inserted: {sorted(section_break_paras)}")

    # ── Step 7: Page breaks — intentionally none (chapters flow continuously) ──
    print(f"[7] Page breaks: 0 (chapters flow continuously)")

    # ── Step 8: Format body paragraphs ──
    for i, p in enumerate(doc.paragraphs):
        # Preserve cover paragraphs untouched (title/abstract/keywords already formatted)
        if i < cover_end:
            continue
        # Also skip abstract body and keywords (formatted in Step 4)
        if abstract_body_idx is not None and i == abstract_body_idx:
            continue
        if keywords_idx is not None and i == keywords_idx:
            continue

        text = p.text.strip()
        has_img = any(r._element.findall(qn('w:drawing')) or
                      r._element.findall(qn('w:pict')) for r in p.runs)

        # Headings
        if i in heading_map:
            h = heading_map[i]
            p.style = doc.styles[h]
            sz_map = {'Heading 1': 15, 'Heading 2': 14, 'Heading 3': 12}
            for run in p.runs:
                set_run_xml(run, sz_map[h], bold=True)
            continue

        # Empty paragraphs (skip)
        if not text and not has_img:
            continue

        # Images
        if has_img:
            p.style = doc.styles['Normal']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            continue

        # Figure/Table captions
        is_cap = False
        if re.match(r'^图\d+[.\-]\d+', text):
            is_cap = True
        elif text.startswith('表') and len(text) < 80 and any(c.isdigit() for c in text):
            is_cap = True

        if is_cap:
            p.style = doc.styles['Normal']
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            for run in p.runs:
                set_run_xml(run, 10.5, bold=True)
            continue

        # References section
        if ref_idx is not None and i > ref_idx:
            p.style = doc.styles['List Paragraph']
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(0)
            for run in p.runs:
                set_run_xml(run, 10.5, bold=False)
            continue

        # Normal body text
        p.style = doc.styles['Normal']
        for run in p.runs:
            set_run_xml(run, 12, bold=False)

    # ── Step 7: Format tables ──
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.first_line_indent = Pt(0)
                    for run in p.runs:
                        set_run_xml(run, 10.5, bold=False)
    print(f"[8] Formatting: {len(heading_map)} headings, {len(doc.tables)} tables")

    # ── Step 9: Save → reload → configure headers/footers ──
    # (Save and reload so python-docx picks up sections created by XML manipulation)

    # Save to temp file
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    os.close(tmp_fd)
    doc.save(tmp_path)
    print(f"[9] Saved to temp: {tmp_path}")

    # Reload
    doc2 = Document(tmp_path)

    # Set page dimensions for all sections
    for sec in doc2.sections:
        sec.page_width = Emu(PAGE_W_EMU)
        sec.page_height = Emu(PAGE_H_EMU)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

    num_sections = len(doc2.sections)
    print(f"[10] Reloaded: {num_sections} sections detected, mode={mode}")

    # Configure headers/footers per section (mode-dependent)
    page_number_restarted = False  # Track whether Arabic page numbers have been restarted

    for i, sec in enumerate(doc2.sections):
        # Get section info from plan (fallback if plan has fewer entries)
        if i < len(section_plan):
            info = section_plan[i]
        else:
            info = {'title': course_title or doc_title, 'section_type': 'chapter'}

        if mode == 'homework':
            if info.get('section_type') == 'abstract':
                # Section 0: Cover + Abstract — NO page numbers, NO header
                clear_header_footer(sec)
                print(f"  Section {i} (abstract): no header, no page numbers")
            else:
                # Section 1+: Body / References — header + page numbers
                set_section_header(sec, course_title or doc_title, size_pt=10.5)
                set_section_footer_page_number(sec, size_pt=10.5)
                if not page_number_restarted:
                    set_page_number_restart(sec, 1)
                    page_number_restarted = True
                print(f"  Section {i} ({info.get('section_type')}): header='{course_title or doc_title}', footer=page#")

        elif mode == 'xsyu_thesis':
            stype = info.get('section_type', 'chapter')

            if stype == 'abstract':
                # Abstract section: NO page numbers, NO header at all
                clear_header_footer(sec)
                print(f"  Section {i} (abstract): no header, no page numbers")

            elif stype == 'toc':
                # TOC section: Roman numeral page numbers, odd/even headers
                enable_even_odd_headers(sec)
                set_section_header(sec, '目录', size_pt=10.5)                   # odd pages
                set_section_even_header(sec, '西安石油大学本科毕业设计(论文)', size_pt=10.5)  # even
                set_section_footer_page_number(sec, size_pt=10.5)
                set_page_number_format(sec, 'upperRoman', start=1)
                print(f"  Section {i} (TOC): Roman numerals from I, odd='目录', even='西安石油大学本科毕业设计(论文)'")

            else:
                # Body chapter section: Arabic page numbers, odd/even headers
                enable_even_odd_headers(sec)
                chapter_title = info['title']
                set_section_header(sec, chapter_title, size_pt=10.5)             # odd pages
                set_section_even_header(sec, '西安石油大学本科毕业设计(论文)', size_pt=10.5)  # even
                set_section_footer_page_number(sec, size_pt=10.5)
                if not page_number_restarted:
                    # First body chapter: restart Arabic page numbering at 1
                    set_page_number_format(sec, 'decimal', start=1)
                    page_number_restarted = True
                else:
                    set_page_number_format(sec, 'decimal')
                print(f"  Section {i} ({stype}): odd='{chapter_title}', even='西安石油大学本科毕业设计(论文)'")

    # ── Step 9: Save final ──
    doc2.save(output_path)

    # Clean up temp
    try:
        os.unlink(tmp_path)
    except Exception:
        pass

    print(f"\n{'='*60}")
    print(f"DONE: {output_path}")
    print(f"Mode: {mode}")
    print(f"Sections: {num_sections}")
    if mode == 'xsyu_thesis':
        print(f"  Section 0 (abstract): no page numbers, no header")
        print(f"  Section 1 (TOC): Roman numerals (Ⅰ,Ⅱ,Ⅲ...), odd='目录', even='西安石油大学本科毕业设计(论文)'")
        print(f"  Sections 2+ (chapters): Arabic page numbers from 1, odd/even headers")
        print(f"  NOTE: After formatting, open in Word and insert TOC (References → Table of Contents)")
        print(f"        at the TOC placeholder between abstract and first chapter.")
    else:
        print(f"  Section 0 (abstract): no page numbers, no header")
        print(f"  Section 1 (body): header='{course_title or doc_title}', page numbers from 1")
        print(f"  Section 2 (references): header='{course_title or doc_title}', continuous page numbers")
    if not no_toc:
        print(f"TOC: Insert in Word → References → Table of Contents → Auto Table")
    print(f"{'='*60}")


# ── CLI entry point ─────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description='Thesis Formatter — 西安石油大学本科毕业设计(论文)模板',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Formal XSYU thesis (default)
  python thesis_formatter.py paper.docx
  python thesis_formatter.py paper.docx output.docx --doc-title="毕业论文"

  # Daily homework mode
  python thesis_formatter.py paper.docx --mode=homework --course-title="高级财务管理"
        """)
    parser.add_argument('input', nargs='?', default=None,
                        help='Input .docx file path')
    parser.add_argument('output', nargs='?', default=None,
                        help='Output .docx file path (default: input_已格式化.docx)')
    parser.add_argument('--doc-title', default='论文',
                        help='Header text / document title (default: 论文)')
    parser.add_argument('--no-toc', action='store_true',
                        help='Skip TOC insertion message')
    parser.add_argument('--mode', default='xsyu_thesis',
                        choices=['xsyu_thesis', 'homework'],
                        help='Formatting mode: xsyu_thesis (XSYU formal thesis) '
                             'or homework (daily assignment)')
    parser.add_argument('--course-title', default=None,
                        help='Course name for header (homework mode only)')

    args = parser.parse_args()

    # Default paths (backward compatibility)
    if args.input is None:
        args.input = "E:/edge/药明康德财务分析报告_已格式化.docx"
    if args.output is None:
        base = args.input.replace('.docx', '')
        args.output = f"{base}_已格式化.docx"

    if not os.path.exists(args.input):
        print(f"ERROR: Input file not found: {args.input}")
        sys.exit(1)

    format_thesis(args.input, args.output, args.doc_title, args.no_toc,
                  mode=args.mode, course_title=args.course_title)


if __name__ == "__main__":
    main()
